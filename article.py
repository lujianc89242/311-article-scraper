import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import re

# Helper function
# Strip the string using regular expressions and return the count of char in that string
def check_sum(str):
    striped_str = str.replace(" ","")
    striped_str = re.sub(r"[\n\t\s]*", "", striped_str)
    return len(striped_str)

def match_rate(output, plain):
    output_char_num = float(check_sum(output))
    plain_char_num = float(check_sum(plain))
    ratio = output_char_num / plain_char_num
    return ratio



r = requests.get("https://portal.311.nyc.gov/all-articles/")
c = r.content

soup = BeautifulSoup(c, "html.parser")
#print(soup.prettify())  # parser the source code and format the html with indentations

############## Collect All URLs for All Articles ###############
outFile = open("article-list.txt","w")
ul = soup.find_all("ul",{"class":"all-az-kas"})
li = ul[0].find_all("li")
url_ctr = 0
url_list = []
for item in li:
    a = item.find_all("a", href=True)
    for item in a:
        #print(item.text)
        href_tag = item['href']
        article_id = href_tag[21:]
        output = article_id + "\n"
        outFile.write(output)
        url_ctr = url_ctr + 1
        full_url = "https://portal.311.nyc.gov/article/?kanumber=" + article_id
        url_list.append(full_url)
outFile.close()

############## Stats output ###############
stats_output = open("stats.txt","w")
err_log_1 = open("err_log_1.txt","w")
err_log_2 = open("err_log_2.txt","w")
log_file = open("log.txt", "w")
article_num = url_ctr
stats_correct_num = 0
stats_total_sum = 0

############## Scraping Information From One Page ###############
# Prepare the scraping object
for each_url in url_list[1302:1303]:
    url = each_url
    article_id = url[45:]
    article_req = requests.get(url)
    article_content = article_req.content
    article_soup = BeautifulSoup(article_content, "html.parser")


    # Get the title of the page
    article_title = article_soup.find_all("h1", {"class":"entry-title page-title"})
    article_title_text = article_title[0].text
    article_title_text_striped = article_title_text
    article_title_text_striped = re.sub(r"[\\\/]*", "", article_title_text_striped)
    print("************************************************************************")
    print("************", article_id, " ", article_title_text, " ******************")
    print("************************************************************************")

    # Get the article content in the page
    ka_content = article_soup.find_all("div", {"class":"ka-content"})[0]

    # Output as .docx file with title as file name
    file_name_docx = "./article/" + article_title_text_striped + "(" + article_id + ").docx"
    document = Document()
    article_heading = str(article_id) + " " + article_title[0].text
    document.add_heading(article_heading,0)

    # Remove all inline CSS style
    for style in ka_content("style"):
        style.decompose()

    # Get only plain text
    plain_text = ka_content.get_text().strip()

    # Start to output the formated article
    output_text = ""  # output_text file has all the output text besides hyperlinks
    cards = ka_content.find_all("div", {"class":"card"})

    for card in cards:
        section_tags2 = []
        card_body = card.find("div",{"class":"card-body"})
        card_head = card.find("button")

        if(str(type(card_head)) == "<class 'bs4.element.Tag'>"):
            print(card_head.text.strip())
            output_text = output_text + card_head.text.strip()
            document.add_heading(card_head.text.strip(), level=1)

        for section in card_body.contents:
            if(str(type(section)) == "<class 'bs4.element.Tag'>"):
                section_tags2.append(section)

        for section_tag in section_tags2:
            print("Section **********")
            for section in section_tag.contents:

                print(section)
                # if(section.name == "p"):
                #     for element in section.contents:
                #         if(str(type(element)) == "<class 'bs4.element.Tag'>"):
                #             if(element.name == "a"):
                #                 url_p = element.text + "(" + element['href'] + ")"
                #                 document.add_paragraph(url_p)
                #                 output_text = output_text + element.text.strip()
                #                 print(url_p)
                #             if(element.name == "strong"):
                #                 document.add_heading(element.text.strip(), level=3)
                #                 output_text = output_text + element.text.strip()
                #                 print(element.text, "(a bold text)")
                #             if(element.name == "br"):
                #                 print("")
                #                 # print("\n")
                #         if(str(type(element)) == "<class 'bs4.element.NavigableString'>" and str(element)!=""):
                #             document.add_paragraph(element)
                #             output_text = output_text + element
                #             print(element)
                # if(section.name == "ul"):
                #     for li in section.contents:
                #         if(str(type(li)) != "<class 'bs4.element.Tag'>" and str(li)!=""):
                #             document.add_paragraph(li)
                #
                #         else:
                #             document.add_paragraph(li.text.strip(), style='List Bullet')
                #             output_text = output_text + li.text.strip()
                #             print(li.text)
                #
                # # Special case when a Need something else? dialog box appears
                # if(section.name == "div"):
                #     document.add_paragraph(section.get_text())
                #     output_text = output_text + section.get_text()
                #
                # if(section.name == "h5"):
                #     try:
                #         value = section.find('input').get('value')
                #         temp_output = value + " (link to a complaint form)"
                #         document.add_paragraph(temp_output)
                #         # output_text = output_text + value
                #         print(temp_output)
                #     except:
                #         pass
                # else:
                #     print("exceptions: ", section )


    # STATS
    print("Plain text: ", check_sum(plain_text), type(plain_text))
    print("output text: ", check_sum(output_text), type(output_text))
    stats_rate = match_rate(output_text,plain_text)
    print(stats_rate)
    stats_total_sum = stats_total_sum + stats_rate
    if stats_rate > 1.0:
        err_log_1_output = article_title_text + ":  " + str(stats_rate) + "\n"
        err_log_1.write(err_log_1_output)
    if stats_rate < 1.0:
        err_log_2_output = article_title_text + ":  " + str(stats_rate) + "\n"
        err_log_2.write(err_log_2_output)

    if stats_rate == 1.0:
        stats_correct_num = stats_correct_num + 1

    document.save(file_name_docx)
    # outFile_txt.close()

temp_output = "Number of articles: " + str(article_num) + "\n"
stats_output.write(temp_output)
temp_output = "Number of articles passed check sum: " + str(stats_correct_num) + "   " + str(stats_correct_num/article_num) + "\n"
stats_output.write(temp_output)
temp_output = "Correct rate: " + str(stats_total_sum/article_num) + "\n"
stats_output.write(temp_output)
err_log_1.close()
err_log_2.close()
stats_output.close()
