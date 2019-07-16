import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches

r = requests.get("https://portal.311.nyc.gov/all-articles/")
c = r.content

soup = BeautifulSoup(c, "html.parser")
#print(soup.prettify())  # parser the source code and format the html with indentations

############## Collect All URLs for All Articles ###############
outFile = open("article-list.txt","w")
ul = soup.find_all("ul",{"class":"all-az-kas"})
li = ul[0].find_all("li")
url_ctr = 0
url_list = []
for item in li:
    a = item.find_all("a", href=True)
    for item in a:
        #print(item.text)
        href_tag = item['href']
        article_id = href_tag[21:]
        output = article_id + "\n"
        #outFile.write(output)
        url_ctr = url_ctr + 1
        full_url = "https://portal.311.nyc.gov/article/?kanumber=" + article_id
        url_list.append(full_url)
# for url in url_list:
#     print(url)
outFile.close()

############## Scraping Information From One Page ###############
# Prepare the scraping object
for each_url in url_list:
    url = each_url
    article_req = requests.get(url)
    article_content = article_req.content
    article_soup = BeautifulSoup(article_content, "html.parser")

    # Get the title of the page
    article_title = article_soup.find_all("h1", {"class":"entry-title page-title"})

    # Get the article content in the page
    ka_content = article_soup.find_all("div", {"class":"ka-content"})[0]

    # Ouput as .txt file
    file_name_txt = "./txt/" + article_title[0].text + ".txt"
    outFile_txt = open(file_name_txt,"w")

    # Ouput as .doc file
    file_name_doc = "./Articles(doc)/" + article_title[0].text + ".doc"
    outFile_doc = open(file_name_doc,"w")

    # Output as .docx file with title as file name
    file_name_docx = "./Articles(docx)/" + article_title[0].text + ".docx"
    document = Document()
    document.add_heading(article_title[0].text,0)

    # Remove all inline CSS style
    for style in ka_content("style"):
        style.decompose()

    # Debugging the ka_content
    # print(ka_content.prettify())

    # Get only plain text
    plain_text = ka_content.get_text().strip()
    for sentence in plain_text.split('.'):
        if sentence == "":
            break
        sentence = sentence.strip()
        sentence = sentence + '.' + '\n'
        plain_text_formated = ""
        for word in sentence.split(" "):
            need_to_split = False
            split_index = 0
            for i in range(len(word)):
                if (word[i].isupper()) and (i != 0):
                    need_to_split = True
                    split_index = i
            # for char in word[1:]:
            #     if char.isupper():
            #         need_to_split = True
            if(need_to_split == False):
                plain_text_formated = plain_text_formated + word + " "
            else:
                plain_text_formated = plain_text_formated + word[:split_index] + "\n" + word[split_index:] + " "

            # for char in word[1:]:
            #     if char.isupper():
            #         print(word, char)
            #         word




        outFile_txt.write(plain_text_formated)


    cards = ka_content.find_all("div", {"class":"card"})
    # Only 1 card case:
    # if(len(cards)==1):
    #     for section in ka_content.find_all("div", {"class":"panel-expand expand"})[1:]:
    #         print(section.get_text())
    # print(len(cards))
    # for card in ka_content.find_all("div", {"clas":"card"}):



    # loop through all descendants recursively for first card
    # Nicer formating but may lose Information
    # for child in ka_content.descendants:
    #     # only capture non-empty NavigableString
    #     # print(child, "         ------------          ", child.name)
    #     if (str(type(child)) == "<class 'bs4.element.NavigableString'>" and str(child.name) != "style"):
    #         print(child.string, "   --------  ", type(child.string), "  ----- ", child.name)
    #         output = child.string + '\n'
    #         outFile.write(output)
        # if(str(child.name) == "a"):
        #     url_p = child.text + "(" + child['href'] + ")"
        #     print(url_p)
        #     document.add_paragraph(url_p)
        # elif(str(child.name) == "button"):
        #     print(child.text.strip())
        #     document.add_heading(child.text.strip(), level=1)
        # elif(str(child.name) == "strong"):
        #     print(child.text.strip())
        #     document.add_heading(child.text.strip(), level=3)
        # elif(str(child.name) == "li"):
        #     print(child.text.strip())
        #     document.add_paragraph(child.text.strip(), style='List Bullet')
        # elif(str(child.name) == "b"):
        #     print(child.text.strip())
        #     p = document.add_paragraph(child.text.strip())
        #     p.bold = True
        # elif(str(child.name) == "p" and child.contents != []):
        #
        #     if(str(child.contents[0].name) == "strong" and (len(child.contents)==1)):
        #         print("do nothing")
        #     if(str(child.contents[0].name) == "em"):
        #         print("do nothing")
        #     elif(str(child.contents[0].name) == "a"):
        #         print("do nothing")
        #     elif(str(child.contents[0].name) == "span"):
        #         print("do nothing")
            # else:
            # print(child.text.strip())
            # document.add_paragraph(child.text.strip())




    # for child in ka_content.descendants:
    #     # only capture non-empty NavigableString
    #     # print(child, "         ------------          ", child.name)
    #     if(str(child.name) == "a"):
    #         url_p = child.text + "(" + child['href'] + ")"
    #         print(url_p)
    #         document.add_paragraph(url_p)
    #     if(str(child.name) == "button"):
    #         print(child.text.strip())
    #         document.add_heading(child.text.strip(), level=1)
    #     if(str(child.name) == "li"):
    #         print(child.text.strip())
    #         document.add_paragraph(child.text.strip(), style='List Bullet')
    #     if(str(child.name) == "p" and child.contents != []):
    #         if(str(child.contents[0].name) == "a"):
    #             print("do nothing")
    #         elif(str(child.contents[0].name) == "span"):
    #             print("do nothing")
    #         else:
    #             print(child.text.strip())
    #             document.add_paragraph(child.text.strip())



    document.save(file_name_docx)
    outFile_txt.close()
